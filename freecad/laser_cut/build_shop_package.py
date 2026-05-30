"""Bundle the laser-cut spec into shop-ready deliverables.

Produces, in `freecad/laser_cut/exports/`:
- `mica_gate_pillar_shop.pdf`  — single PDF, page 1 = Vietnamese spec sheet,
  page 2 = the cut layout (merged from `mica_gate_pillar.pdf`).
- `mica_gate_pillar_spec.docx` — Word doc of the same spec sheet, so the
  shop can copy parameters or annotate.

Run after `build_mica_pillar.py`:
    python -m freecad.laser_cut.build_shop_package
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
)
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


SPEC_TITLE = "MICA GATE PILLAR — LASER CUT"
SPEC_SUBTITLE = "Smart gate pillar enclosure (Pi 4 + ESP32 + sensors)"
SHEET_W_MM = 1000
SHEET_H_MM = 600
MICA_THICKNESS_MM = 5

PIECES = [
    ("1. FRONT",  "150 × 228 mm",   "Mặt trước (không cutout)"),
    ("2. BACK",   "150 × 300 mm",   "Mặt sau, lỗ adapter Ø10mm + 6 lỗ M3"),
    ("3. LEFT",   "240 × 300 mm",   "Pentagon, mặt tháo, 4 lỗ M3 corner + 4 lỗ M3 PCB"),
    ("4. RIGHT",  "240 × 300 mm",   "Pentagon, khe arm 20×120mm + HC-SR04 cluster"),
    ("5. TOP",    "150 × 168 mm",   "Mặt trên, LCD 98×40 + camera + USB pass-through 20×12"),
    ("6. SLOPE",  "150 × 102 mm",   "Mặt vát 45° (trống, RFID đọc xuyên qua)"),
    ("7. BOTTOM", "150 × 240 mm",   "Mặt đáy, 5 khe thông gió 30×3mm"),
    ("8. ARM",    "115 × 15 mm",    "Cần chắn, 3 sọc engrave 20×5mm (sơn ĐỎ sau khi cắt)"),
]


def _try_register_vietnamese_font():
    """Try to register a system Vietnamese-capable TrueType font."""
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            name = Path(path).stem.replace("-", "")
            pdfmetrics.registerFont(TTFont(name, path))
            return name
    return None


def _build_spec_pdf(out_path: Path) -> None:
    """Render the spec sheet as a single A4 PDF page."""
    font_name = _try_register_vietnamese_font() or "Helvetica"

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=SPEC_TITLE,
    )

    body = ParagraphStyle(
        "body", fontName=font_name, fontSize=10, leading=13, alignment=TA_LEFT,
    )
    h1 = ParagraphStyle(
        "h1", fontName=font_name, fontSize=18, leading=22, spaceAfter=4,
    )
    h2 = ParagraphStyle(
        "h2", fontName=font_name, fontSize=11, leading=14, spaceBefore=10, spaceAfter=4,
    )

    story = []
    story.append(Paragraph(SPEC_TITLE, h1))
    story.append(Paragraph(SPEC_SUBTITLE, body))
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>Vật liệu &amp; khổ tấm</b>", h2))
    story.append(Paragraph(
        f"Mica acrylic <b>trong suốt</b>, độ dày <b>{MICA_THICKNESS_MM} mm</b>. "
        f"Khổ tấm: <b>{SHEET_W_MM} × {SHEET_H_MM} mm</b> (1 tấm đủ 8 chi tiết, "
        f"tổng diện tích cắt ~0.30 m²).",
        body,
    ))

    story.append(Paragraph("<b>Quy ước trên file</b>", h2))
    story.append(Paragraph(
        "• Đường <font color='red'><b>ĐỎ (RGB 255,0,0)</b></font>: <b>CẮT THỦNG</b><br/>"
        "• Đường <b>ĐEN (RGB 0,0,0)</b>: <b>KHẮC NÔNG</b> (engrave, sâu ~0.2mm, không cắt thủng)<br/>"
        "• Stroke 0.3/0.5mm chỉ để hiển thị PDF — máy laser cắt theo đường tâm geometry, "
        "không theo độ dày stroke.",
        body,
    ))

    story.append(Paragraph("<b>Danh sách 8 chi tiết</b>", h2))
    table_data = [["#", "Mảnh", "Kích thước", "Ghi chú"]]
    for piece, dim, note in PIECES:
        num, name = piece.split(". ", 1)
        table_data.append([num, name, dim, note])
    tbl = Table(table_data, colWidths=[10 * mm, 24 * mm, 30 * mm, 110 * mm])
    tbl.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, -1), font_name, 9),
        ("FONT", (0, 0), (-1, 0), font_name, 9),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dddddd")),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(tbl)

    story.append(Paragraph("<b>Dung sai</b>", h2))
    story.append(Paragraph(
        "±0.2 mm cho lỗ và mép. Đường cắt cùng tấm cách nhau ≥5 mm để tránh vỡ mica.",
        body,
    ))

    story.append(Paragraph("<b>Lưu ý ARM</b>", h2))
    story.append(Paragraph(
        "Sau khi cắt, sơn ĐỎ vào 3 vạch khắc trên thanh ARM để tạo sọc đỏ-trắng "
        "giống barrier thật. 2 lỗ Ø2.2mm gần đầu là điểm bắt vít M2 vào servo horn.",
        body,
    ))

    story.append(Spacer(1, 8))
    story.append(Paragraph(
        "<i>Trang sau: layout cắt (8 mảnh trên khổ {} × {} mm).</i>"
        .format(SHEET_W_MM, SHEET_H_MM),
        body,
    ))

    doc.build(story)


def _build_spec_docx(out_path: Path) -> None:
    """Write the spec as a Word .docx for shop use."""
    doc = Document()

    # Title
    title = doc.add_heading(SPEC_TITLE, level=1)
    doc.add_paragraph(SPEC_SUBTITLE)

    doc.add_heading("Vật liệu & khổ tấm", level=2)
    p = doc.add_paragraph()
    p.add_run("Mica acrylic trong suốt, độ dày ").font.size = Pt(11)
    p.add_run(f"{MICA_THICKNESS_MM} mm").bold = True
    p.add_run(". Khổ tấm: ")
    p.add_run(f"{SHEET_W_MM} × {SHEET_H_MM} mm").bold = True
    p.add_run(" (1 tấm đủ 8 chi tiết, tổng diện tích cắt ~0.30 m²).")

    doc.add_heading("Quy ước trên file", level=2)
    doc.add_paragraph("Đường ĐỎ (RGB 255,0,0): CẮT THỦNG", style="List Bullet")
    doc.add_paragraph(
        "Đường ĐEN (RGB 0,0,0): KHẮC NÔNG (engrave ~0.2mm, không thủng)",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Stroke 0.3/0.5mm chỉ để hiển thị PDF — máy laser cắt theo đường tâm geometry.",
        style="List Bullet",
    )

    doc.add_heading("Danh sách 8 chi tiết", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Mảnh"
    hdr[2].text = "Kích thước"
    hdr[3].text = "Ghi chú"
    for piece, dim, note in PIECES:
        num, name = piece.split(". ", 1)
        row = table.add_row().cells
        row[0].text = num
        row[1].text = name
        row[2].text = dim
        row[3].text = note

    doc.add_heading("Dung sai", level=2)
    doc.add_paragraph(
        "±0.2 mm cho lỗ và mép. Đường cắt cùng tấm cách nhau ≥5 mm để tránh vỡ mica."
    )

    doc.add_heading("Lưu ý ARM", level=2)
    doc.add_paragraph(
        "Sau khi cắt, sơn ĐỎ vào 3 vạch khắc trên thanh ARM để tạo sọc đỏ-trắng "
        "giống barrier thật. 2 lỗ Ø2.2mm gần đầu là điểm bắt vít M2 vào servo horn."
    )

    doc.add_paragraph()
    note = doc.add_paragraph(
        f"File cắt: mica_gate_pillar.pdf (hoặc .svg). "
        f"File này: spec sheet đính kèm."
    )
    note.runs[0].italic = True

    doc.save(out_path)


def main():
    out_dir = Path(__file__).parent / "exports"
    out_dir.mkdir(parents=True, exist_ok=True)

    cut_pdf = out_dir / "mica_gate_pillar.pdf"
    if not cut_pdf.exists():
        raise SystemExit(
            f"{cut_pdf} not found — run `python -m freecad.laser_cut.build_mica_pillar` first."
        )

    spec_pdf = out_dir / "mica_gate_pillar_spec.pdf"
    shop_pdf = out_dir / "mica_gate_pillar_shop.pdf"
    spec_docx = out_dir / "mica_gate_pillar_spec.docx"

    _build_spec_pdf(spec_pdf)
    _build_spec_docx(spec_docx)

    # Merge spec + cut into one shop-ready PDF.
    writer = PdfWriter()
    writer.append(str(spec_pdf))
    writer.append(str(cut_pdf))
    with open(shop_pdf, "wb") as fh:
        writer.write(fh)

    print(f"Wrote {spec_pdf}")
    print(f"Wrote {spec_docx}")
    print(f"Wrote {shop_pdf}  (2 trang: spec + layout cắt — file để gửi shop)")


if __name__ == "__main__":
    main()
